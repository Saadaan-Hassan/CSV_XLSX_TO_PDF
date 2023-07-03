# Importing libraries
import sys
import csv
import pandas as pd
import os
import requests
import openpyxl

from io import BytesIO
from datetime import date
from docx2pdf import convert

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor

# Detect the delimiter of the csv file
def detect_csv_delimiter(file_path):
    with open(file_path, "r", newline="", encoding='utf-8', errors='ignore') as file:
        sample = file.read(4096)  # Read a sample of the file

        dialect = csv.Sniffer().sniff(sample)
        delimiter = dialect.delimiter

    return delimiter


# Read the csv file
def read_csv(file, delimiter):
    df = pd.read_csv(file, delimiter=delimiter, encoding='utf-8')

    return df


# Extract Characteristics from the csv file
def extract_characteristics(att_col, primary_tokens):
    # Create a dictionary for the characteristics of all tokens
    all_characteristics = {}

    for primary_token in primary_tokens:
        # Get the rows with the same ProductPrimaryToken
        specific_rows = att_col[att_col["ProductPrimaryToken"] == primary_token]
        specific_df = pd.DataFrame(specific_rows)

        # Create a dictionary with the characteristics of the specific ProductPrimaryToken
        characteristics = {}

        # Iterate over the columns in the specific DataFrame
        for column in specific_df.columns:
            if column != "ProductPrimaryToken":  # Exclude the "ProductPrimaryToken" column
                # Check the number of unique values in the column
                num_unique_values = specific_df[column].nunique()

                if num_unique_values == 1:
                    value = specific_df[column].iloc[0]
                    if pd.notna(value) and value != "" and value != "None":
                        characteristics[column] = value

        all_characteristics[primary_token] = characteristics

    return all_characteristics


# Extract Attributes of a row
def extract_attributes(row, characteristics):
    attributes = {}

    for key, value in row.items():
        if key not in characteristics and pd.notna(value) and value != "" and value != "None":
            attributes[key] = value

    return attributes

# Create a word document
def create_word_document(data):
    # Create the pdf file path
    pdf_file = data["ProductPrimaryToken"] + "_" + data["Token"] + ".pdf"
    file_path = "PDF/" + str(date.today()) + "/" + pdf_file
    try:
        if not os.path.exists(file_path):
            # Create a new Word document
            doc = Document()

            # Set page size (e.g., A4)
            section = doc.sections[0]
            section.page_width = Inches(8.27)  # Width of A4 in inches
            section.page_height = Inches(11.69)  # Height of A4 in inches

            # Set margins (e.g., 1 inch on all sides)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

            # Set the default paragraph style
            default_style = doc.styles['Normal']
            default_style.font.size = Pt(11)  # Set default font size to 11 points
            default_style.font.name = 'Calibri'  # Set default font to Calibri
            default_style.paragraph_format.line_spacing = 1.5  # Set line spacing to 1.5 times the font size

            # Calculate the table and paragraph widths based on page dimensions
            table_width = int(section.page_width - section.left_margin - section.right_margin)
            text_column_width = int(table_width * 0.5)  # Adjust as desired
            image_column_width = int(table_width - text_column_width)

            # Add a table with 2 columns
            table = doc.add_table(rows=1, cols=2)
            table.allow_autofit = False
            table.columns[0].width = image_column_width
            table.columns[1].width = text_column_width

            # Get the first row of the table
            row = table.rows[0]

            if isinstance(data["Image_ProductPrimary"], str):
    # Add an image from a URL on the left side with custom size
                image_urls = data["Image_ProductPrimary"].split(",")
            else:
                image_urls = str(data["Image_ProductPrimary"]).split(",")

            if len(image_urls) > 0:
                if image_urls[0] != "None" and image_urls[0] != "nan":
                    if requests.head(image_urls[0]).status_code == 200:
                        print("Downloading Primary " + image_urls[0] + "...")
                        response = requests.get(image_urls[0])
                        image_data = BytesIO(response.content)
                    else:
                        print("No image found for " + data["ProductPrimaryToken"] + "!")
                        image_data = ('image-not-found.png')
                elif len(image_urls) > 1 and image_urls[1] != "None" and image_urls[1] != "nan":
                    if requests.head(image_urls[1]).status_code == 200:
                        print("Downloading Secondary " + image_urls[1] + "...")
                        response = requests.get(image_urls[1])
                        image_data = BytesIO(response.content)
                    else:
                        print("No image found for " + data["ProductPrimaryToken"] + "!")
                        image_data = ('image-not-found.png')
                else:
                    print("No image found for " + data["ProductPrimaryToken"] + "!")
                    image_data = ('image-not-found.png')
            else:
                print("No image URLs found for " + data["ProductPrimaryToken"] + "!")
                image_data = ('image-not-found.png')
            
            cell_image = row.cells[0]
            cell_image.width = Inches(3.7)
            cell_image.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Align content vertically to center
            cell_image.add_paragraph().add_run().add_picture(image_data, width=Inches(3.61))

            # Add text on the right side
            cell_text = row.cells[1]
            cell_text.width = Inches(3.5)
            cell_text.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Align content vertically to center
            cell_text_paragraph = cell_text.add_paragraph()

            # Add the text content
            cell_text_paragraph.add_run("Referencia de Producto: ").bold = True
            cell_text_paragraph.add_run(data["ProductPrimaryToken"] + "\n")
            cell_text_paragraph.add_run("Nombre de Producto: ").bold = True
            cell_text_paragraph.add_run(data["Name_es"] + "\n")
            cell_text_paragraph.add_run("Descripción de Producto:\n").bold = True
            cell_text_paragraph.add_run(str(data["ProductSection_T2_INFO_es"]).replace("\n", "") + "\n")

            # Add a heading with black font color
            heading = doc.add_heading('Bullet', level=1)
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set RGB values to (0, 0, 0) for black


            # Split the bullet points by the new line character
            if isinstance(data["Attribute_BulletPointsProducto"], str):
                if "\r\n" in data["Attribute_BulletPointsProducto"]:
                    data["Attribute_BulletPointsProducto"] = data[
                        "Attribute_BulletPointsProducto"
                    ].split("\r\n")
                elif "\n" in data["Attribute_BulletPointsProducto"]:
                    data["Attribute_BulletPointsProducto"] = data[
                        "Attribute_BulletPointsProducto"
                    ].split("\n")
                else:
                    data["Attribute_BulletPointsProducto"] = [
                        data["Attribute_BulletPointsProducto"]
                    ]
                data["Attribute_BulletPointsProducto"] = [
                    item.lstrip("- ") for item in data["Attribute_BulletPointsProducto"]
                ]
            elif isinstance(data["Attribute_BulletPointsProducto"], float):
                data["Attribute_BulletPointsProducto"] = str(
                    data["Attribute_BulletPointsProducto"]
                ).split("\r\n")
                data["Attribute_BulletPointsProducto"] = [
                    item.lstrip("- ") for item in data["Attribute_BulletPointsProducto"] if item
                ]


            # Add bullet points under the heading in 2 columns
            bullet_points = data["Attribute_BulletPointsProducto"]
            # Create a paragraph with bullet points
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(15)  # Adjust the value as needed for top padding
            paragraph_format.space_after = Pt(15)   # Adjust the value as needed for bottom padding

            for bullet_point in bullet_points:
                paragraph.add_run('• ').bold = True  # Add bullet symbol (you can customize it)
                paragraph.add_run(bullet_point + '\n')

            # Apply shading (background color fill) to the paragraph
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d6d6d6"/>')
            paragraph._element.get_or_add_pPr().append(shading)

            # Add a heading for characteristics
            heading = doc.add_heading('Caracteristicas', level=1)
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set RGB values to (0, 0, 0) for black


            # Create a list of characteristics
            characteristics = []

            # Iterate over the characteristics dictionaries to create the list
            for key, val in data["Characteristics"].items():
                key = key.replace("Attribute_", "")
                characteristics.append(f"{key}: {val}")

            # Determine the number of rows needed for characteristics table
            num_rows_characteristics = (len(characteristics) + 1) // 2

            # Add a new table for characteristics bullet points
            characteristics_table = doc.add_table(rows=num_rows_characteristics, cols=2)
            characteristics_table.allow_autofit = False
            characteristics_table.columns[0].width = Inches(3.5)
            characteristics_table.columns[1].width = Inches(3.5)

            # Set table properties for background shading
            tbl_props = characteristics_table._element.xpath('.//w:tblPr')
            tbl_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d6d6d6"/>')
            tbl_props.append(tbl_shading)


            # Populate the bullet points in the characteristics table
            row_index = 0
            col_index = 0
            counter = 0

            for i in range(len(characteristics)):
                cell = characteristics_table.cell(row_index, col_index)
                cell.width = Inches(3.5)

                if counter < len(characteristics):
                    cell.text = characteristics[counter]

                counter += 1

                col_index += 1
                if col_index >= 2:
                    col_index = 0
                    row_index += 1

            # Set cell shading (background color) for cells under "Caracteristicas" heading
            for row in characteristics_table.rows:
                for cell in row.cells:
                    if cell.text:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Adjust vertical alignment
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d6d6d6"/>')
                        cell._element.tcPr.append(shading)
                        paragraph = cell.paragraphs[0]
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_before = Pt(5)  # Adjust the value as needed for top padding
                        paragraph_format.space_after = Pt(5)   # Adjust the value as needed for bottom padding

            # Add a heading
            heading = doc.add_heading('Atributos', level=1)
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set RGB values to (0, 0, 0) for black


            # Add a new table for attributes bullet points
            attr_table = doc.add_table(rows=2, cols=2)
            attr_table.allow_autofit = False
            attr_table.columns[0].width = Inches(3.5)
            attr_table.columns[1].width = Inches(3.5)

            # Create a list of attributes
            attributes = []

            # Iterate over the attributes dictionaries to create the list
            for key, val in data["Attributes"].items():
                key = key.replace("Attribute_", "")
                attributes.append(f"{key}: {val}")

           # Populate the table with attributes
            counter = 0
            for attribute in attributes:
                row_index = counter // 2
                col_index = counter % 2

                if col_index == 0:
                    attr_table.add_row()

                cell = attr_table.cell(row_index, col_index)
                cell.width = Inches(3.5)
                cell.text = attribute

                counter += 1


            # Set cell shading (background color) for cells under "Caracteristicas" heading
            for row in attr_table.rows:
                for cell in row.cells:
                    if cell.text:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Adjust vertical alignment
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d6d6d6"/>')
                        cell._element.tcPr.append(shading)
                        paragraph = cell.paragraphs[0]
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_before = Pt(5)  # Adjust the value as needed for top padding
                        paragraph_format.space_after = Pt(5)   # Adjust the value as needed for bottom padding


            # Save the Word document 
            file = data["ProductPrimaryToken"] + "_" + data["Token"]
            doc.save("temp.docx")
            
            create_pdf(file + ".pdf")

    except KeyboardInterrupt:
        # Actions to be taken if the program is interrupted
        print("Program interrupted. Cleaning up...")
        os.remove("temp.docx")
# Create a pdf file of the row and save it in the PDF folder of the current date folder
def create_pdf(pdf_file):

    # Create the pdf file of the row and save it in the PDF folder of the current date folder
    if not os.path.exists("PDF"):
        print("Creating PDF folder...")
        os.makedirs("PDF")
    if not os.path.exists("PDF/" + str(date.today())):
        print("Creating " + str(date.today()) + " folder...")
        os.makedirs("PDF/" + str(date.today()))

    # Create the pdf file path
    file_path = "PDF/" + str(date.today()) + "/" + pdf_file

    print("Creating " + pdf_file + "...")
    convert("temp.docx", file_path)

# Main function to read the csv or xlsx file and extract the characteristics and attributes of each row
def main(file_path):
    # Get the file extension
    _, file_extension = os.path.splitext(file_path)

    # Read the csv or xlsx file
    if file_extension == ".csv":
        print("Reading " + file_path + "...")
        df = read_csv(file_path, detect_csv_delimiter(file_path)).sort_values(
            by="ProductPrimaryToken"
        )

    elif file_extension == ".xlsx":
        print("Reading " + file_path + "...")
        # Load the Excel file
        wb = openpyxl.load_workbook(file_path)
        sheet_name = wb.sheetnames[0]  # Assuming you want to read the first sheet
        sheet = wb[sheet_name]
        data = sheet.values
        df = pd.DataFrame(data)

        # Set column names if needed
        df.columns = df.iloc[0]

        # Convert all columns to string
        df = df.astype(str)

        # Sort DataFrame by 'ProductPrimaryToken' column
        df = df.sort_values(by="ProductPrimaryToken")

    # Drop columns with all NaN values and columns with all 0 values
    df = df.dropna(axis=1, how="all").drop(columns=df.columns[df.eq(0.0).all()])
        
    # Drop rows with all NaN values
    df = df.dropna(axis=0, how="all")

    # Reset the index
    df = df.reset_index(drop=True)


    # Get columns 2, 3, and 4
    columns = df.iloc[:, [1, 2, 3, 4, 6, 8]]

    # Select columns that start with "Attribute_" (excluding "Attribute_BulletPointsProducto" and "Attribute_Estado")
    attribute_columns = df.filter(
        regex="^(?!Attribute_BulletPointsProducto|Attribute_Estado)(Attribute_)"
    ).copy()  # Create a copy of the filtered DataFrame

    # Add the ProductPrimaryToken column to the attribute_columns DataFrame
    attribute_columns["ProductPrimaryToken"] = df["ProductPrimaryToken"]

    # Get duplicate attribute columns based on ProductPrimaryToken
    duplicate_attribute_columns = attribute_columns[
        attribute_columns.duplicated(subset="ProductPrimaryToken", keep=False)
    ]

    # Get the unique tokens
    unique_tokens = columns["ProductPrimaryToken"].unique()

    # Create a dictionary to store the data of each row
    data = {
        "Token": "",
        "ProductPrimaryToken": "",
        "Name_es": "",
        "ProductSection_T2_INFO_es": "",
        "Image_ProductPrimary": "",
        "Attribute_BulletPointsProducto": "",
        "Characteristics": {},
        "Attributes": {},
    }

    print("Extracting characteristics for every PrimaryProductToken...")
    # Extract the characteristics and attributes of each row and create a pdf file
    chars = extract_characteristics(duplicate_attribute_columns, unique_tokens)

    # Iterate over the rows in the DataFrame and create a pdf file for each row
    for i in range(len(columns["ProductPrimaryToken"])):
        char = chars[
            columns["ProductPrimaryToken"][i]
        ]  # Get the characteristics of the row

        # Extract the attributes of the row
        attr = extract_attributes(attribute_columns.iloc[i].to_dict(), char)
        attr.pop("ProductPrimaryToken", None)

        # Add the data of the row to the data dictionary
        data["Token"] = columns["Token"][i]
        data["ProductPrimaryToken"] = columns["ProductPrimaryToken"][i]
        data["Name_es"] = columns["Name_es"][i]
        data["ProductSection_T2_INFO_es"] = columns["ProductSection_T2_INFO_es"][i]
        data["Image_ProductPrimary"] = columns["Image_ProductPrimary"][i]
        data["Attribute_BulletPointsProducto"] = columns[
            "Attribute_BulletPointsProducto"
        ][i]
        data["Characteristics"] = char
        data["Attributes"] = attr

        create_word_document(data)


if __name__ == "__main__":
    try:
        # Check if the file path is provided as an argument
        if len(sys.argv) > 1:
            file_path = sys.argv[1]
            main(file_path)
            os.remove("temp.docx")
        else:
            print("Please provide the file path as an argument.")
    
    except KeyboardInterrupt:
        # Actions to be taken if the program is interrupted
        print("Program interrupted. Cleaning up...")
        os.remove("temp.docx")
